from fastapi import Depends, HTTPException, Query
from typing import Annotated, Dict, Optional
from app_fastapi.configurations.configuration import get_settings
from app_fastapi.schemas.ocr.request_http import OCRType
from app_fastapi.utilities.ocr_util.v1.ocr_util import call_API
from fastapi import UploadFile, File
from typing import List
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO


async def http_post(ocr: OCRType = Query(...)):
    """
    미리 업로드 된 이미지 파일 OCR로 전체 추출 작업
    """
    if ocr == "google":
        print("구글")

    ocr_api = call_API(ocr)

    source_path = "storage/ocr/source"
    target_path = "storage/ocr/target"

    for filename in os.listdir("storage/ocr/source"):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff')):
            file_path = os.path.join(source_path,filename)

            name, ext = os.path.splitext(filename)
            save_path = os.path.join(target_path, ocr, name + ".docx")
            print(filename) 
            ocr_result = ocr_api.run_process(
                target = file_path
            )
            final_text_lst = OCRResult2List(
                    ocr_result = ocr_result,
                    file_name = ""
            )
            
            List2Docx(
                input_lst = final_text_lst,
                save_path = save_path
            )

    return True


def OCRResult2List(ocr_result: List, file_name:str) -> List:
    ocr_result_lst = []
    for page_num, data in reorder_with_page(ocr_result).items():
        one_line_text = make_1line(data)

        # file_name_flag = f"<<file_name_flag>><<{os.path.basename(file_name)}>>"
        # page_num_flag = f"<<page_num_flag>><<{str(page_num).zfill(4)}>>"

        # ocr_result_lst.append(file_name_flag)
        # ocr_result_lst.append(page_num_flag)
        # ocr_result_lst.append("\n")
        for line in one_line_text:
            ocr_result_lst.append(line)

    return ocr_result_lst


def reorder_with_page(input_data:List) -> dict:
    page_dict = {}
    for i in input_data:
        page = i['page']
        coords = i['coords']
        text = i['text']

        if page not in page_dict:
            page_dict[page] = []
        
        page_dict[page].append((coords, text))

    return page_dict


def make_1line(result:list, cv_img = None) -> list: # bbox 그린 이미지가 필요하면 cv_img 받기
    # make tmp_dict
    tmp_dict = {}     
    for idx, data in enumerate(result):
        coors = data[0]
        text = data[1]
        
        y_lst = []
        
        for coor in coors:
            coor['x'] = int(coor['x'])
            coor['y'] = int(coor['y'])
            
            y_lst.append(coor['y'])
            
        # cv_img = cv2.rectangle(cv_img, coors[0], coors[2], (0, 255, 0), 2) # bbox 그린 이미지가 필요하면 주석해제
            
        height = max(y_lst) - min(y_lst)
        y_center = int(min(y_lst) + (height/2))
        
        data_dict = {
            'height': height,
            'y_center': y_center,
            'text': text
        }
                
        tmp_dict[idx] = data_dict
        
    # make exist_next_lst
    exist_next_lst = []

    for i in range(len(tmp_dict)-1):
        target_height = tmp_dict[i]['height']
        target_y_center = tmp_dict[i]['y_center']

        next_y_center = tmp_dict[i+1]['y_center']

        exist_next = target_y_center - int(target_height/4) <= next_y_center <= target_y_center + int(target_height/4)

        if exist_next:
            exist_next_lst.append(i)
    exist_next_lst = sorted(exist_next_lst)
    
    # make final_index_lst
    origin_index_lst = sorted([i for i in range(len(tmp_dict))])

    final_index_lst = []
    tmp_index_lst = []

    for integ in origin_index_lst:
        if integ not in exist_next_lst: # 다음에 이어지는게 없는 경우
            final_index_lst.append([integ])

        else: # 다음에 이어지는게 있는 경우
            if integ + 1 not in exist_next_lst: # 다음에 이어지는게 1개인 경우
                if len(tmp_index_lst) == 0: # 그동안 쌓인게 없는 경우
                    final_index_lst.append([integ, integ+1])
                    origin_index_lst.remove(integ+1)
                else: # 그동안 쌓인게 있는 경우
                    tmp_index_lst.append(integ)
                    tmp_index_lst.append(integ+1)
                    tmp_index_lst = sorted(tmp_index_lst)
                    final_index_lst.append(tmp_index_lst)
                    origin_index_lst.remove(integ+1)
                    tmp_index_lst = []
            else: # 다음에 이어지는게 2개 이상인 경우
                tmp_index_lst.append(integ)
                
    # make final_text_lst
    final_text_lst = []
    for tmp_i in final_index_lst:
        tmp_str = ''
        for j in tmp_i:
            tmp_str += tmp_dict[j]['text'] + ' '
        tmp_str = tmp_str[:-1]
        final_text_lst.append(tmp_str)

    return final_text_lst # bbox 그린 이미지가 필요하면 cv_img retrun하기


def List2Docx(input_lst: List, save_path: str):
    doc = Document()
    for final_text in input_lst:
        # 기존코드 
        # doc.add_paragraph(final_text)
        
        ### 241223 글씨체 지정 추가 ###
        paragraph = doc.add_paragraph()
        set_font_style(paragraph, final_text)

    # doc을 메모리에 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # doc을 local에 저장
    doc.save(save_path)


def set_font_style(paragraph, text):
    """
    텍스트에 따라 글씨체를 지정하는 함수.
    - 한글: 맑은 고딕
    - 그 외: Arial
    """
    buffer = ""
    current_lang = None

    for char in text:
        # 한글 여부 확인
        is_korean = '\uAC00' <= char <= '\uD7A3' or '\u1100' <= char <= '\u11FF' or '\u3131' <= char <= '\u318E'
        lang = "korean" if is_korean else "other"

        # 언어가 변경되었으면 이전 버퍼를 처리하고 초기화
        if current_lang is not None and current_lang != lang:
            run = paragraph.add_run(buffer)
            if current_lang == "korean":
                run.font.name = "맑은 고딕"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            else:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            buffer = ""

        # 현재 문자 추가
        buffer += char
        current_lang = lang

    # 마지막 버퍼 처리
    if buffer:
        run = paragraph.add_run(buffer)
        if current_lang == "korean":
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        else:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
